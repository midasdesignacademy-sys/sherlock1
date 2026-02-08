"""
SHERLOCK - Document Ingestion Agent (Agent 1)
Soul: docs/agents/agent_1_soul.md
Orchestrates document ingestion, OCR, normalization, deduplication, quarantine, cryptography_findings.
"""

import hashlib
import re
import shutil
import time
import unicodedata
import uuid
from pathlib import Path
from typing import Dict, List, Any, Tuple
from datetime import datetime

from loguru import logger

from core.state import InvestigationState, DocumentMetadata
from core.config import settings
from core.persistence import (
    get_doc_status,
    log_doc_start,
    log_doc_success,
    log_doc_failed,
    STATUS_DONE,
)

try:
    import magic
except ImportError:
    magic = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import fitz
except ImportError:
    fitz = None

try:
    import langdetect
except ImportError:
    langdetect = None

import email
from email import policy


def _detect_file_type(file_path: Path) -> str:
    if magic:
        try:
            mime = magic.Magic(mime=True)
            return mime.from_file(str(file_path))
        except Exception:
            pass
    return file_path.suffix.lower() or "application/octet-stream"


def _calculate_hash(file_path: Path) -> str:
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def _normalize_text(text: str) -> str:
    if not text:
        return ""
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"[\x00-\x1f\x7f-\x9f]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _detect_language(text: str) -> str:
    if not text or not langdetect:
        return "unknown"
    try:
        return langdetect.detect(text) or "unknown"
    except Exception:
        return "unknown"


def _quarantine_file(file_path: Path, reason: str) -> None:
    try:
        settings.QUARANTINE_DIR.mkdir(parents=True, exist_ok=True)
        dest = settings.QUARANTINE_DIR / f"{file_path.stem}_{uuid.uuid4().hex[:8]}{file_path.suffix}"
        shutil.copy2(file_path, dest)
        logger.warning(f"Quarantined {file_path.name} to {dest} ({reason})")
    except Exception as e:
        logger.error(f"Quarantine failed for {file_path.name}: {e}")


class DocumentIngestionAgent:
    """Agent 1: Ingest documents, extract text, normalize, deduplicate. Soul-aligned."""

    def __init__(self):
        self.supported = set(s.lower() for s in settings.SUPPORTED_FORMATS)
        self.max_size = settings.MAX_FILE_SIZE_MB * 1024 * 1024

    def process(self, state: InvestigationState) -> InvestigationState:
        logger.info("[Agent 1] Starting document ingestion...")
        try:
            upload_path = state.get("config", {}).get("uploads_path")
            if upload_path:
                upload_dir = Path(upload_path)
            else:
                upload_dir = settings.UPLOADS_DIR

            if not upload_dir.exists():
                state["error_log"] = state.get("error_log", []) + [f"Upload dir not found: {upload_dir}"]
                return state

            user_descriptions: Dict[str, str] = {}
            desc_path = Path(upload_dir) / "descriptions.json"
            if desc_path.exists():
                try:
                    import json
                    data = json.loads(desc_path.read_text(encoding="utf-8"))
                    if isinstance(data, dict):
                        user_descriptions = {k: str(v) for k, v in data.items() if v and k != "descriptions.json"}
                except Exception:
                    pass

            files = [f for f in upload_dir.iterdir() if f.is_file() and f.name != "descriptions.json"]
            if not files:
                logger.warning("No files found in uploads directory")
                state["error_log"] = state.get("error_log", []) + ["No files to ingest"]
                return state

            existing_hashes = set()
            for meta in state.get("document_metadata", {}).values():
                if isinstance(meta, dict):
                    existing_hashes.add(meta.get("file_hash"))
                else:
                    existing_hashes.add(getattr(meta, "file_hash", None))

            processed_docs = list(state.get("processed_docs", []))
            document_metadata = dict(state.get("document_metadata", {}))
            extracted_text = dict(state.get("extracted_text", {}))
            raw_documents = list(state.get("raw_documents", []))
            cryptography_findings = list(state.get("cryptography_findings", []))
            investigation_id = (state.get("config") or {}).get("investigation_id") or ""

            for file_path in files:
                if file_path.suffix.lower() not in self.supported:
                    logger.warning(f"Unsupported format: {file_path.name}")
                    continue
                if file_path.stat().st_size > self.max_size:
                    logger.warning(f"File too large (>{settings.MAX_FILE_SIZE_MB}MB): {file_path.name}")
                    continue

                file_hash = _calculate_hash(file_path)
                if file_hash in existing_hashes:
                    logger.info(f"Skipping duplicate: {file_path.name}")
                    continue
                if get_doc_status(file_hash, investigation_id) == STATUS_DONE:
                    logger.info(f"Skipping already processed (ledger DONE): {file_path.name}")
                    continue
                existing_hashes.add(file_hash)

                log_doc_start(file_hash, investigation_id)
                try:
                    file_type = _detect_file_type(file_path)
                    doc_id = file_hash[:16]
                    t0 = time.perf_counter()

                    text_content, status, extraction_method, ocr_confidence, page_count, meta_extra, crypto_finding = self._process_one(
                        file_path, file_type, doc_id
                    )

                    processing_time_ms = int((time.perf_counter() - t0) * 1000)

                    if status == "encrypted" and crypto_finding:
                        cryptography_findings.append(crypto_finding)

                    if status == "failed":
                        _quarantine_file(file_path, meta_extra.get("error_message", "extraction failed"))

                    if not text_content and status not in ("encrypted", "failed"):
                        status = "partial" if status == "success" else status
                        text_content = ""

                    text_content = _normalize_text(text_content) if text_content else ""
                    language = _detect_language(text_content) if text_content else "unknown"

                    author, created, modified = self._extract_metadata_dates(file_path, file_type, file_path.suffix.lower())
                    meta_dict = meta_extra or {}
                    meta_dict.update({
                        "author": author,
                        "creation_date": created.isoformat() if created else None,
                        "modification_date": modified.isoformat() if modified else None,
                        "title": None,
                        "producer": None,
                    })
                    user_desc = user_descriptions.get(file_path.name, "").strip()
                    if user_desc:
                        meta_dict["user_description"] = user_desc

                    meta = DocumentMetadata(
                        doc_id=doc_id,
                        filename=file_path.name,
                        file_type=file_type,
                        file_hash=file_hash,
                        size_bytes=file_path.stat().st_size,
                        upload_timestamp=datetime.now(),
                        source=str(file_path.parent),
                        language=language,
                        author=author,
                        created=created,
                        modified=modified,
                        file_path=str(file_path),
                        status=status,
                        extraction_method=extraction_method,
                        ocr_confidence=ocr_confidence,
                        processing_time_ms=processing_time_ms,
                        page_count=page_count,
                        error_message=meta_extra.get("error_message") if isinstance(meta_extra, dict) else None,
                        metadata=meta_dict if isinstance(meta_dict, dict) else None,
                    )
                    document_metadata[doc_id] = meta.model_dump()
                    extracted_text[doc_id] = text_content
                    processed_docs.append({"doc_id": doc_id, "text": text_content, "metadata": meta.model_dump()})
                    raw_documents.append({"doc_id": doc_id, "file_path": str(file_path), "metadata": meta.model_dump()})
                    if status == "success" or status == "partial":
                        log_doc_success(file_hash, investigation_id)
                        logger.info(f"Ingested: {file_path.name} ({len(text_content)} chars, {extraction_method})")
                    else:
                        log_doc_failed(file_hash, investigation_id, "ingest_documents")
                        logger.warning(f"Document {file_path.name}: status={status}")
                except Exception as e:
                    log_doc_failed(file_hash, investigation_id, "ingest_documents")
                    logger.exception(f"Document {file_path.name}: {e}")
                    state["error_log"] = state.get("error_log", []) + [f"Ingestion doc error {file_path.name}: {str(e)}"]

            state["processed_docs"] = processed_docs
            state["document_metadata"] = document_metadata
            state["extracted_text"] = extracted_text
            state["raw_documents"] = raw_documents
            state["cryptography_findings"] = cryptography_findings
            state["current_step"] = "ingestion_complete"
            logger.info(f"[Agent 1] Ingested {len(document_metadata)} documents")
        except Exception as e:
            logger.exception(f"[Agent 1] Error: {e}")
            state["error_log"] = state.get("error_log", []) + [f"Ingestion error: {str(e)}"]
        return state

    def _process_one(
        self, file_path: Path, file_type: str, doc_id: str
    ) -> Tuple[str, str, str, float, int, Any, Any]:
        status = "success"
        extraction_method = "unknown"
        ocr_confidence = 0.0
        page_count = 0
        meta_extra = {}
        crypto_finding = None
        text_content = ""

        suffix = file_path.suffix.lower()
        try:
            if "pdf" in file_type or suffix == ".pdf":
                text_content, status, extraction_method, ocr_confidence, page_count, meta_extra, crypto_finding = self._extract_pdf(file_path, doc_id)
                return (text_content, status, extraction_method, ocr_confidence, page_count, meta_extra, crypto_finding)
            if "word" in file_type or suffix in [".docx", ".doc"]:
                text_content = self._extract_docx(file_path)
                extraction_method = "python-docx"
                return (text_content, status, extraction_method, ocr_confidence, page_count, meta_extra, crypto_finding)
            if "excel" in file_type or suffix in [".xlsx", ".xls"]:
                text_content = self._extract_excel(file_path)
                extraction_method = "pandas_excel"
                return (text_content, status, extraction_method, ocr_confidence, page_count, meta_extra, crypto_finding)
            if "text" in file_type or suffix == ".txt":
                text_content = file_path.read_text(encoding="utf-8", errors="ignore")
                extraction_method = "plaintext"
                return (text_content, status, extraction_method, ocr_confidence, page_count, meta_extra, crypto_finding)
            if suffix == ".csv":
                text_content = self._extract_csv(file_path)
                extraction_method = "pandas_csv"
                return (text_content, status, extraction_method, ocr_confidence, page_count, meta_extra, crypto_finding)
            if suffix == ".eml" or "email" in (file_type or ""):
                text_content = self._extract_eml(file_path)
                extraction_method = "email"
                return (text_content, status, extraction_method, ocr_confidence, page_count, meta_extra, crypto_finding)
            if suffix in [".png", ".jpg", ".jpeg"] or (file_type and "image" in file_type):
                text_content, ocr_confidence = self._extract_image(file_path)
                extraction_method = "tesseract_ocr" if ocr_confidence else "image_skip"
                return (text_content, status, extraction_method, ocr_confidence, page_count, meta_extra, crypto_finding)
            try:
                from unstructured.partition.auto import partition
                elements = partition(filename=str(file_path))
                text_content = "\n".join([getattr(el, "text", str(el)) for el in elements if getattr(el, "text", None)])
                extraction_method = "unstructured"
                return (text_content, status, extraction_method, ocr_confidence, page_count, meta_extra, crypto_finding)
            except Exception as e:
                meta_extra = {"error_message": str(e)}
                status = "unsupported"
                return ("", status, "unstructured_failed", ocr_confidence, page_count, meta_extra, crypto_finding)
        except Exception as e:
            logger.error(f"Extract failed {file_path.name}: {e}")
            return ("", "failed", extraction_method or "error", ocr_confidence, page_count, {"error_message": str(e)}, crypto_finding)

    def _extract_pdf(self, file_path: Path, doc_id: str) -> Tuple[str, str, str, float, int, Dict, Any]:
        status = "success"
        extraction_method = "pymupdf"
        ocr_confidence = 0.0
        page_count = 0
        meta_extra = {}
        crypto_finding = None
        text_parts = []

        if fitz:
            try:
                doc = fitz.open(file_path)
                if getattr(doc, "is_encrypted", False):
                    doc.close()
                    crypto_finding = {
                        "document_id": doc_id,
                        "finding_type": "pdf_encrypted",
                        "requires_password": True,
                        "confidence": 1.0,
                    }
                    return ("", "encrypted", "pymupdf_encrypted", 0.0, 0, {"error_message": "PDF is password-protected"}, crypto_finding)
                page_count = len(doc)
                for pnum in range(page_count):
                    page = doc.load_page(pnum)
                    t = page.get_text() or ""
                    if t.strip():
                        text_parts.append(t)
                doc.close()
                text = "\n\n".join(text_parts)
                if len(text.strip()) >= 50:
                    return (text, status, extraction_method, ocr_confidence, page_count, meta_extra, crypto_finding)
            except Exception as e:
                if "encrypted" in str(e).lower() or "password" in str(e).lower():
                    crypto_finding = {"document_id": doc_id, "finding_type": "pdf_encrypted", "requires_password": True, "confidence": 1.0}
                    return ("", "encrypted", "pymupdf_encrypted", 0.0, 0, {"error_message": str(e)}, crypto_finding)
                if "data" in str(e).lower() or "FileDataError" in type(e).__name__:
                    if pdfplumber:
                        try:
                            with pdfplumber.open(file_path, repair=True) as pdf:
                                text_parts = []
                                for page in pdf.pages:
                                    t = (page.extract_text() or "") if page else ""
                                    if t.strip():
                                        text_parts.append(t)
                                text = "\n\n".join(text_parts)
                                page_count = len(pdf.pages)
                                if text.strip():
                                    return (text, status, "pdfplumber_repair", ocr_confidence, page_count, meta_extra, crypto_finding)
                        except Exception as e2:
                            meta_extra = {"error_message": f"pdfplumber repair failed: {e2}"}
                            return ("", "failed", "pdfplumber_repair", 0.0, 0, meta_extra, crypto_finding)
                    meta_extra = {"error_message": str(e)}
                    return ("", "failed", "pymupdf", 0.0, 0, meta_extra, crypto_finding)
                meta_extra = {"error_message": str(e)}
                return ("", "failed", "pymupdf", 0.0, 0, meta_extra, crypto_finding)

        if pypdf:
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                page_count = len(reader.pages)
                for page in reader.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        text_parts.append(t)
            text = "\n\n".join(text_parts)
            if len(text.strip()) >= 50:
                return (text, status, "pypdf", ocr_confidence, page_count, meta_extra, crypto_finding)

        if pytesseract and fitz and Image:
            try:
                import io
                doc = fitz.open(file_path)
                if getattr(doc, "is_encrypted", False):
                    doc.close()
                    crypto_finding = {"document_id": doc_id, "finding_type": "pdf_encrypted", "requires_password": True, "confidence": 1.0}
                    return ("", "encrypted", "tesseract_ocr", 0.0, 0, {"error_message": "PDF is encrypted"}, crypto_finding)
                page_count = len(doc)
                ocr_parts = []
                for pnum in range(page_count):
                    page = doc.load_page(pnum)
                    pix = page.get_pixmap(dpi=150)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    if getattr(settings, "TESSERACT_CMD", None):
                        pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
                    ocr_text = pytesseract.image_to_string(img, lang=getattr(settings, "OCR_LANGUAGES", "por+eng"))
                    if ocr_text.strip():
                        ocr_parts.append(ocr_text)
                doc.close()
                if ocr_parts:
                    text = "\n\n".join(ocr_parts)
                    ocr_confidence = 0.85
                    return (text, status, "tesseract_ocr", ocr_confidence, page_count, meta_extra, crypto_finding)
            except Exception as e:
                logger.debug(f"PDF OCR failed: {e}")
                meta_extra = {"error_message": str(e)}

        text = "\n\n".join(text_parts)
        return (text, status if text.strip() else "partial", extraction_method, ocr_confidence, page_count, meta_extra, crypto_finding)

    def _extract_docx(self, file_path: Path) -> str:
        if not DocxDocument:
            return ""
        doc = DocxDocument(file_path)
        return "\n\n".join(p.text for p in doc.paragraphs if p.text)

    def _extract_excel(self, file_path: Path) -> str:
        if not pd:
            return ""
        xl = pd.ExcelFile(file_path)
        parts = []
        for name in xl.sheet_names:
            df = pd.read_excel(file_path, sheet_name=name)
            parts.append(f"=== {name} ===\n{df.to_string()}")
        return "\n\n".join(parts)

    def _extract_csv(self, file_path: Path) -> str:
        if not pd:
            return file_path.read_text(encoding="utf-8", errors="ignore")
        df = pd.read_csv(file_path)
        return df.to_string()

    def _extract_eml(self, file_path: Path) -> str:
        raw = file_path.read_bytes()
        try:
            msg = email.message_from_bytes(raw, policy=policy.default)
            parts = []
            if msg.get("subject"):
                parts.append(f"Subject: {msg['subject']}")
            if msg.get("from"):
                parts.append(f"From: {msg['from']}")
            if msg.get("to"):
                parts.append(f"To: {msg['to']}")
            if msg.get("date"):
                parts.append(f"Date: {msg['date']}")
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    ct = part.get_content_type()
                    if ct == "text/plain":
                        body = part.get_content()
                        break
                    if ct == "text/html" and not body:
                        body = part.get_content()
                        if body and re.search(r"<[^>]+>", str(body)):
                            body = re.sub(r"<[^>]+>", " ", str(body))
                        break
            else:
                body = msg.get_content()
            if body:
                parts.append(str(body))
            return "\n\n".join(parts) if parts else ""
        except Exception as e:
            logger.error(f"EML parse failed: {e}")
            return file_path.read_text(encoding="utf-8", errors="ignore")

    def _extract_image(self, file_path: Path) -> Tuple[str, float]:
        if not pytesseract or not Image:
            return "", 0.0
        try:
            img = Image.open(file_path)
            if getattr(settings, "TESSERACT_CMD", None):
                pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
            text = pytesseract.image_to_string(img, lang=getattr(settings, "OCR_LANGUAGES", "por+eng"))
            return (text or "", 0.85)
        except Exception as e:
            logger.error(f"Image OCR failed {file_path.name}: {e}")
            return ("", 0.0)

    def _extract_metadata_dates(self, file_path: Path, file_type: str, suffix: str) -> Tuple[Any, Any, Any]:
        author, created, modified = None, None, None
        try:
            if suffix == ".pdf" and pypdf:
                with open(file_path, "rb") as f:
                    reader = pypdf.PdfReader(f)
                    meta = reader.metadata
                    if meta:
                        author = meta.get("/Author") or meta.get("Author")
            if suffix in [".docx", ".doc"] and DocxDocument:
                doc = DocxDocument(file_path)
                core_props = getattr(doc, "core_properties", None)
                if core_props:
                    author = getattr(core_props, "author", None) or author
                    created = getattr(core_props, "created", None)
                    modified = getattr(core_props, "modified", None)
        except Exception as e:
            logger.debug(f"Metadata extraction failed: {e}")
        return (author, created, modified)
